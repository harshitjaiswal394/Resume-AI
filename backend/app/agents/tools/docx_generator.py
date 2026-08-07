"""
DOCX Generator — builds an ATS-friendly, downloadable Word document from a
tailored resume (the JSON produced by the Resume Tailor agent).

Uses python-docx (already in requirements). Single-column layout, standard
fonts, no tables/graphics so ATS parsers read it cleanly.
"""

from __future__ import annotations

import io
import logging
from typing import Any, Dict, List

logger = logging.getLogger("resumatch-ai.tools.docx")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None  # type: ignore


def _style_document(doc: Any) -> None:
    """Apply clean ATS-friendly base styling."""
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(0)


def _add_section_header(doc: Any, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    # Thin divider under the section header
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F2937")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bullet(doc: Any, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)


def _clean_text(value: Any) -> str:
    """Coerce any value to a single-line string, dropping control chars."""
    if value is None:
        return ""
    text = str(value).strip()
    return " ".join(text.split())


def build_tailored_docx(
    tailored: Dict[str, Any],
    *,
    filename: str = "tailored-resume.docx",
) -> bytes:
    """
    Generate a DOCX from the tailor agent's `tailored_resume` JSON.

    Returns the raw `.docx` file bytes. Raises RuntimeError if python-docx
    is not installed.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for DOCX generation")

    doc = Document()
    _style_document(doc)

    full_name = _clean_text(tailored.get("fullName"))
    if full_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(full_name)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x11, 0x16, 0x1E)
        p.paragraph_format.space_after = Pt(6)

    contact = _clean_text(tailored.get("contact"))
    if not contact:
        contact_parts = []
        email = _clean_text(tailored.get("email"))
        if email:
            contact_parts.append(email)
        phone = _clean_text(tailored.get("phone"))
        if phone:
            contact_parts.append(phone)
        links = tailored.get("links") or {}
        if isinstance(links, dict):
            for key in ("github", "linkedin", "portfolio"):
                val = _clean_text(links.get(key))
                if val:
                    contact_parts.append(val)
        contact = "  |  ".join(contact_parts)
    if contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(contact)
        p.paragraph_format.space_after = Pt(10)

    summary = _clean_text(tailored.get("summary"))
    if summary:
        _add_section_header(doc, "Professional Summary")
        doc.add_paragraph(summary)

    skills = tailored.get("skills") or []
    if skills:
        _add_section_header(doc, "Skills")
        if isinstance(skills, list):
            doc.add_paragraph(", ".join(_clean_text(s) for s in skills))
        else:
            doc.add_paragraph(_clean_text(skills))

    experience = tailored.get("experience") or []
    if experience:
        _add_section_header(doc, "Professional Experience")
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            title = _clean_text(exp.get("title"))
            company = _clean_text(exp.get("company"))
            duration = _clean_text(exp.get("duration"))
            location = _clean_text(exp.get("location"))

            header_parts = [p for p in [title, company] if p]
            if header_parts:
                p = doc.add_paragraph()
                run = p.add_run(" | ".join(header_parts))
                run.bold = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(1)

            meta_parts = [p for p in [location, duration] if p]
            if meta_parts:
                p = doc.add_paragraph()
                run = p.add_run(" — ".join(meta_parts))
                run.italic = True
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                p.paragraph_format.space_after = Pt(3)

            for bullet in exp.get("bullets", []) or []:
                if isinstance(bullet, dict):
                    text = _clean_text(bullet.get("text") or bullet.get("original_bullet"))
                else:
                    text = _clean_text(bullet)
                if text:
                    _add_bullet(doc, text)

    education = tailored.get("education") or []
    if education:
        _add_section_header(doc, "Education")
        for edu in education:
            if isinstance(edu, str):
                doc.add_paragraph(_clean_text(edu))
                continue
            if not isinstance(edu, dict):
                continue
            degree = _clean_text(edu.get("degree"))
            institution = _clean_text(edu.get("institution"))
            year = _clean_text(edu.get("year"))
            parts = [p for p in [degree, institution] if p]
            if parts:
                p = doc.add_paragraph()
                run = p.add_run(" — ".join(parts))
                run.bold = True
                p.paragraph_format.space_after = Pt(0)
            if year:
                p = doc.add_paragraph()
                run = p.add_run(year)
                run.italic = True
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                p.paragraph_format.space_after = Pt(2)

    projects = tailored.get("projects") or []
    if projects:
        _add_section_header(doc, "Projects")
        for proj in projects:
            if isinstance(proj, str):
                doc.add_paragraph(_clean_text(proj))
                continue
            if not isinstance(proj, dict):
                continue
            title = _clean_text(proj.get("title"))
            if title:
                p = doc.add_paragraph()
                run = p.add_run(title)
                run.bold = True
                p.paragraph_format.space_after = Pt(0)
            desc = _clean_text(proj.get("description") or proj.get("text"))
            if desc:
                _add_bullet(doc, desc)

    internships = tailored.get("internships") or []
    if internships:
        _add_section_header(doc, "Internships")
        for intern in internships:
            if isinstance(intern, str):
                _add_bullet(doc, _clean_text(intern))
                continue
            if not isinstance(intern, dict):
                continue
            title = _clean_text(intern.get("title"))
            company = _clean_text(intern.get("company"))
            duration = _clean_text(intern.get("duration"))
            header_parts = [p for p in [title, company] if p]
            if header_parts:
                p = doc.add_paragraph()
                run = p.add_run(" | ".join(header_parts))
                run.bold = True
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)
            if duration:
                p = doc.add_paragraph()
                run = p.add_run(duration)
                run.italic = True
                run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                p.paragraph_format.space_after = Pt(2)
            desc = _clean_text(intern.get("description") or intern.get("text"))
            if desc:
                _add_bullet(doc, desc)

    certifications = tailored.get("certifications") or []
    if certifications:
        _add_section_header(doc, "Certifications")
        for cert in certifications:
            if isinstance(cert, str):
                doc.add_paragraph(_clean_text(cert))
                continue
            if isinstance(cert, dict):
                name = _clean_text(cert.get("name") or cert.get("title"))
                if name:
                    doc.add_paragraph(_clean_text(name))

    languages = tailored.get("languages") or []
    if languages:
        _add_section_header(doc, "Languages")
        items = []
        for lang in languages:
            if isinstance(lang, str):
                items.append(_clean_text(lang))
            elif isinstance(lang, dict):
                name = _clean_text(lang.get("language"))
                prof = _clean_text(lang.get("proficiency"))
                items.append(f"{name} — {prof}" if name and prof else (name or prof))
        if items:
            doc.add_paragraph(", ".join(i for i in items if i))

    achievements = tailored.get("achievements") or []
    if achievements:
        _add_section_header(doc, "Achievements")
        for ach in achievements:
            if isinstance(ach, str):
                _add_bullet(doc, _clean_text(ach))
                continue
            if isinstance(ach, dict):
                title = _clean_text(ach.get("title") or ach.get("name"))
                desc = _clean_text(ach.get("description"))
                if title and desc:
                    _add_bullet(doc, f"{title}: {desc}")
                elif title:
                    _add_bullet(doc, title)
                elif desc:
                    _add_bullet(doc, desc)

    buffer = io.BytesIO()
    doc.save(buffer)
    logger.info("DOCX_GENERATED | bytes=%d", buffer.tell())
    return buffer.getvalue()


def safe_filename(full_name: str = "", version_number: int = 0) -> str:
    """Build a safe, readable download filename."""
    name = full_name.strip().replace(" ", "-") if full_name else "tailored-resume"
    name = "".join(c for c in name if c.isalnum() or c in "-_")
    if version_number > 0:
        return f"{name}-v{version_number}.docx"
    return f"{name}.docx"
